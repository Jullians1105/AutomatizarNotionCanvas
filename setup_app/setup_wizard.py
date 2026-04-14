"""
Canvas → Notion Sync  |  Setup Wizard
Guia interactiva para configurar la aplicacion desde cero.
"""
import os
import sys
import re
import json
import subprocess
import threading
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
from pathlib import Path
import requests


# ── Colores / fuentes ──────────────────────────────────────────────────────────
BG       = "#1e1e2e"
BG2      = "#2a2a3e"
ACCENT   = "#7c3aed"
ACCENT2  = "#a78bfa"
FG       = "#e2e8f0"
FG2      = "#94a3b8"
SUCCESS  = "#22c55e"
ERROR    = "#ef4444"
WARN     = "#f59e0b"
FONT     = ("Segoe UI", 10)
FONT_B   = ("Segoe UI", 10, "bold")
FONT_H   = ("Segoe UI", 14, "bold")
FONT_SM  = ("Segoe UI", 9)


def resource_path(relative: str) -> str:
    """Obtiene ruta correcta dentro del .exe (PyInstaller) o en desarrollo."""
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative)
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), relative)


# ── Validaciones ──────────────────────────────────────────────────────────────
def validate_canvas_token(token: str, base_url: str) -> tuple[bool, str]:
    try:
        r = requests.get(
            f"{base_url.rstrip('/')}/users/self",
            headers={"Authorization": f"Bearer {token}"},
            timeout=10,
        )
        if r.status_code == 200:
            name = r.json().get("name", "usuario")
            return True, f"Conectado como: {name}"
        return False, f"Token inválido (HTTP {r.status_code})"
    except Exception as e:
        return False, f"Error de conexión: {e}"


def validate_notion_token(token: str) -> tuple[bool, str]:
    try:
        r = requests.get(
            "https://api.notion.com/v1/users/me",
            headers={
                "Authorization": f"Bearer {token}",
                "Notion-Version": "2022-06-28",
            },
            timeout=10,
        )
        if r.status_code == 200:
            name = r.json().get("name", "workspace")
            return True, f"Token válido — workspace: {name}"
        return False, f"Token inválido (HTTP {r.status_code})"
    except Exception as e:
        return False, f"Error de conexión: {e}"


def validate_notion_db(token: str, db_id: str) -> tuple[bool, str]:
    clean = db_id.replace("-", "").strip()
    if len(clean) != 32:
        return False, "El ID debe tener 32 caracteres hexadecimales"
    try:
        r = requests.get(
            f"https://api.notion.com/v1/databases/{db_id}",
            headers={
                "Authorization": f"Bearer {token}",
                "Notion-Version": "2022-06-28",
            },
            timeout=10,
        )
        if r.status_code == 200:
            title_arr = r.json().get("title", [])
            title = title_arr[0]["plain_text"] if title_arr else "sin título"
            return True, f"Base de datos encontrada: {title}"
        return False, f"No se pudo acceder (HTTP {r.status_code}) — ¿compartiste la BD con la integración?"
    except Exception as e:
        return False, f"Error de conexión: {e}"


def validate_telegram(bot_token: str, chat_id: str) -> tuple[bool, str]:
    try:
        r = requests.post(
            f"https://api.telegram.org/bot{bot_token}/sendMessage",
            json={"chat_id": chat_id, "text": "✅ Canvas→Notion configurado correctamente!"},
            timeout=10,
        )
        if r.status_code == 200:
            return True, "Mensaje de prueba enviado a Telegram"
        detail = r.json().get("description", "")
        return False, f"Error Telegram: {detail}"
    except Exception as e:
        return False, f"Error de conexión: {e}"


# ── Guardado de .env ──────────────────────────────────────────────────────────
def save_env(dest: Path, values: dict):
    lines = [
        f"CANVAS_API_TOKEN={values['canvas_token']}",
        f"CANVAS_BASE_URL={values['canvas_url']}",
        f"NOTION_API_TOKEN={values['notion_token']}",
        f"NOTION_DATABASE_ID={values['notion_db']}",
        f"TELEGRAM_BOT_TOKEN={values['tg_token']}",
        f"TELEGRAM_CHAT_ID={values['tg_chat']}",
    ]
    dest.write_text("\n".join(lines) + "\n", encoding="utf-8")


# ── Generador del Word ────────────────────────────────────────────────────────
def generate_guide_docx(dest: Path):
    """Genera el Word de instrucciones paso a paso."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return False, "python-docx no está instalado"

    doc = Document()

    # Estilos generales
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_heading(text, level=1, color=(124, 58, 237)):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(*color)
        return p

    def add_para(text, bold=False, color=None, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_step(num, title, body_lines):
        p = doc.add_paragraph()
        run = p.add_run(f"Paso {num}: {title}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(124, 58, 237)
        for line in body_lines:
            doc.add_paragraph(line, style="List Bullet")

    def add_link_note(text):
        p = doc.add_paragraph()
        run = p.add_run(f"🔗 {text}")
        run.italic = True
        run.font.color.rgb = RGBColor(99, 102, 241)

    # ── Portada ───────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Canvas → Notion Sync")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(124, 58, 237)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Guía de configuración paso a paso")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    doc.add_paragraph(
        "Esta guía te ayuda a obtener todas las credenciales necesarias para que "
        "la aplicación sincronice automáticamente tus tareas de Canvas LMS con tu "
        "base de datos de Notion y te notifique por Telegram."
    )
    doc.add_page_break()

    # ── SECCIÓN 1: Canvas ─────────────────────────────────────────────────────
    add_heading("1. Token de Canvas LMS", level=1)
    doc.add_paragraph(
        "El token de Canvas te permite que la app lea tus cursos y tareas "
        "de forma segura sin necesidad de tu contraseña."
    )

    add_step(1, "Inicia sesión en Canvas", [
        "Abre tu navegador y entra a tu campus de Canvas",
        "Ejemplo: https://umb.instructure.com  (o el dominio de tu universidad)",
    ])
    add_link_note("URL de tu Canvas: https://[tu-universidad].instructure.com")

    add_step(2, "Ve a Configuración de cuenta", [
        "Haz clic en tu foto de perfil (esquina superior izquierda)",
        "Selecciona  'Configuración'  o  'Account > Settings'",
    ])

    add_step(3, "Genera el token de acceso", [
        "Baja hasta la sección  'Tokens de acceso aprobados'",
        "Haz clic en  '+ Nuevo token de acceso'",
        "En 'Propósito' escribe:  Canvas Notion Sync",
        "Deja la fecha de expiración vacía (sin expiración) o pon una fecha futura",
        "Haz clic en  'Generar token'",
        "⚠️  COPIA el token ahora — Canvas no te lo mostrará de nuevo",
    ])

    add_para("Ejemplo de token Canvas:", bold=True)
    add_para("  7~AbCdEfGhIjKlMnOpQrStUvWxYz1234567890AbCdEfGhIjKlMnOpQrStUvWx", color=(71, 85, 105))

    doc.add_page_break()

    # ── SECCIÓN 2: Notion Integración ─────────────────────────────────────────
    add_heading("2. Integración y Token de Notion", level=1)
    doc.add_paragraph(
        "Notion usa 'integraciones' (bots internos) para permitir que apps externas "
        "lean y escriban en tus bases de datos."
    )

    add_step(1, "Crea una integración en Notion", [
        "Ve a:  https://www.notion.so/my-integrations",
        "Haz clic en  '+ Nueva integración'",
        "Dale un nombre:  Canvas Sync",
        "Tipo:  Interna (Internal)",
        "Workspace: selecciona tu workspace personal",
        "Haz clic en  'Enviar' o 'Submit'",
    ])
    add_link_note("https://www.notion.so/my-integrations")

    add_step(2, "Copia el token secreto", [
        "Dentro de la integración creada, ve a la pestaña  'Secrets'  o  'Configuración'",
        "Verás  'Token secreto interno'  o  'Internal Integration Secret'",
        "Haz clic en  'Mostrar' y luego copia el token",
        "Empieza con:  ntn_  o  secret_",
    ])

    add_para("Ejemplo de token Notion:", bold=True)
    add_para("  ntn_AbCdEfGhIjKlMnOpQrStUvWxYz1234567890AbCd", color=(71, 85, 105))

    doc.add_page_break()

    # ── SECCIÓN 3: Base de datos Notion ───────────────────────────────────────
    add_heading("3. Base de datos de Notion e ID", level=1)
    doc.add_paragraph(
        "La aplicación necesita una base de datos en Notion donde guardará tus tareas. "
        "Puedes crear una nueva o usar una existente."
    )

    add_step(1, "Crea la base de datos", [
        "Abre Notion y crea una nueva página",
        "Escribe  /base de datos  y selecciona  'Base de datos - completa'",
        "Nómbrala como quieras, ejemplo:  Tareas Canvas",
    ])

    add_step(2, "Agrega las columnas necesarias", [
        "La base de datos debe tener estas propiedades:",
        "  • Descripción  (Título — ya existe por defecto)",
        "  • Materia        (Selección — Select)",
        "  • Estado de tarea (Estado — Status)",
        "  • Fecha Limite   (Fecha — Date)",
        "  • Personas       (Personas — People)",
    ])

    add_step(3, "Conecta la integración a la base de datos", [
        "Abre tu base de datos en Notion",
        "Haz clic en los  '...'  (tres puntos) en la esquina superior derecha",
        "Selecciona  'Conexiones'  o  'Connections'",
        "Busca tu integración  'Canvas Sync'  y haz clic en  'Confirmar'",
        "⚠️  Sin este paso la app no podrá escribir en la BD",
    ])

    add_step(4, "Obtén el ID de la base de datos", [
        "Abre la base de datos en el navegador (no en la app de escritorio)",
        "La URL se verá así:",
        "  https://www.notion.so/TU-WORKSPACE/AbCdEfGh1234...?v=...",
        "El ID es la parte larga de 32 caracteres ANTES del '?v='",
        "Ejemplo:  AbCdEfGh1234AbCdEfGh1234AbCdEfGh",
        "También puede tener guiones:  abcdef12-3456-7890-abcd-ef1234567890",
    ])

    add_para("Ejemplo de ID de base de datos:", bold=True)
    add_para("  abcdef12-3456-7890-abcd-ef1234567890", color=(71, 85, 105))

    doc.add_page_break()

    # ── SECCIÓN 4: Telegram Bot ────────────────────────────────────────────────
    add_heading("4. Bot y Chat ID de Telegram", level=1)
    doc.add_paragraph(
        "La app te enviará un mensaje de Telegram cada vez que sincronice tus tareas. "
        "Necesitas crear un bot y saber tu chat ID."
    )

    add_step(1, "Crea el bot con BotFather", [
        "Abre Telegram y busca:  @BotFather",
        "Envía el comando:  /newbot",
        "BotFather te pedirá:",
        "  • Nombre del bot (ej: Canvas Notion Sync Bot)",
        "  • Username del bot — debe terminar en 'bot' (ej: canvasnotionsync_bot)",
        "BotFather te dará el  TOKEN DEL BOT  — cópialo",
        "Empieza con números seguidos de ':' y letras/números",
    ])
    add_link_note("Busca @BotFather en Telegram")

    add_para("Ejemplo de token de bot:", bold=True)
    add_para("  1234567890:AAFbCdEfGhIjKlMnOpQrStUvWxYz_AbCdEfG", color=(71, 85, 105))

    add_step(2, "Obtén tu Chat ID", [
        "Opción A — Usando @userinfobot:",
        "  • Busca en Telegram:  @userinfobot",
        "  • Envíale cualquier mensaje",
        "  • Te responderá con tu ID de usuario (un número como: 123456789)",
        "",
        "Opción B — Usando la API de Telegram:",
        "  • Envía un mensaje a tu nuevo bot",
        "  • Abre en el navegador:",
        "    https://api.telegram.org/bot<TU_TOKEN>/getUpdates",
        "  • Busca  'chat':{'id': NUMERO}  en la respuesta JSON",
    ])

    add_step(3, "Inicia una conversación con el bot", [
        "Busca tu bot por su username en Telegram",
        "Envíale  /start  para iniciar la conversación",
        "⚠️  El bot no puede enviarte mensajes si nunca le has escrito primero",
    ])

    doc.add_page_break()

    # ── SECCIÓN 5: GitHub Actions ─────────────────────────────────────────────
    add_heading("5. Configurar GitHub Actions (ejecución automática)", level=1)
    doc.add_paragraph(
        "GitHub Actions permite que la sincronización se ejecute automáticamente en la nube, "
        "sin necesidad de tener tu computador encendido. Las credenciales se guardan como "
        "Secrets cifrados en tu repositorio — nunca quedan expuestas en el código."
    )

    add_step(1, "Sube el proyecto a GitHub", [
        "Si aún no tienes el repositorio en GitHub, créalo en:  https://github.com/new",
        "Nombre sugerido:  canvas-notion-sync",
        "Visibilidad:  Privado (recomendado, ya que el .env NO se sube)",
        "Desde la terminal en la carpeta del proyecto ejecuta:",
        "  git remote add origin https://github.com/TU_USUARIO/canvas-notion-sync.git",
        "  git push -u origin main",
    ])
    add_link_note("https://github.com/new")

    add_step(2, "Agrega los Secrets al repositorio", [
        "Ve a tu repositorio en GitHub",
        "Haz clic en  Settings  (la rueda dentada arriba a la derecha)",
        "En el menu izquierdo, expande  Secrets and variables  y haz clic en  Actions",
        "Haz clic en  New repository secret  para cada una de las siguientes variables:",
        "",
        "  Nombre: CANVAS_API_TOKEN        Valor: tu token de Canvas",
        "  Nombre: CANVAS_BASE_URL         Valor: https://umb.instructure.com/api/v1",
        "  Nombre: NOTION_API_TOKEN        Valor: tu token de Notion",
        "  Nombre: NOTION_DATABASE_ID      Valor: el ID de tu base de datos",
        "  Nombre: TELEGRAM_BOT_TOKEN      Valor: el token de tu bot",
        "  Nombre: TELEGRAM_CHAT_ID        Valor: tu chat ID de Telegram",
        "",
        "⚠️  El nombre debe ser EXACTAMENTE igual (mayusculas incluidas)",
    ])

    add_step(3, "Crea el archivo de workflow", [
        "Dentro de tu proyecto crea la carpeta:  .github/workflows/",
        "Dentro de esa carpeta crea el archivo:  sync.yml",
        "Copia y pega el siguiente contenido en ese archivo:",
    ])

    # Bloque de código YAML
    yaml_content = (
        "name: Canvas Notion Sync\n\n"
        "on:\n"
        "  schedule:\n"
        "    # Lunes a viernes: 9am y 9pm hora Colombia (UTC-5 = 14:00 y 02:00 UTC)\n"
        "    - cron: '0 14 * * 1-5'\n"
        "    - cron: '0 2  * * 2-6'\n"
        "    # Sabado: 4pm Colombia (21:00 UTC)\n"
        "    - cron: '0 21 * * 6'\n"
        "  workflow_dispatch:   # Permite ejecutarlo manualmente desde GitHub\n\n"
        "jobs:\n"
        "  sync:\n"
        "    runs-on: ubuntu-latest\n"
        "    steps:\n"
        "      - name: Checkout del repositorio\n"
        "        uses: actions/checkout@v4\n\n"
        "      - name: Configurar Python\n"
        "        uses: actions/setup-python@v5\n"
        "        with:\n"
        "          python-version: '3.11'\n\n"
        "      - name: Instalar dependencias\n"
        "        run: pip install -r requirements.txt\n\n"
        "      - name: Ejecutar sincronizacion\n"
        "        env:\n"
        "          CANVAS_API_TOKEN:   ${{ secrets.CANVAS_API_TOKEN }}\n"
        "          CANVAS_BASE_URL:    ${{ secrets.CANVAS_BASE_URL }}\n"
        "          NOTION_API_TOKEN:   ${{ secrets.NOTION_API_TOKEN }}\n"
        "          NOTION_DATABASE_ID: ${{ secrets.NOTION_DATABASE_ID }}\n"
        "          TELEGRAM_BOT_TOKEN: ${{ secrets.TELEGRAM_BOT_TOKEN }}\n"
        "          TELEGRAM_CHAT_ID:   ${{ secrets.TELEGRAM_CHAT_ID }}\n"
        "        run: python main.py\n"
    )
    p_yaml = doc.add_paragraph()
    run_yaml = p_yaml.add_run(yaml_content)
    run_yaml.font.name = "Courier New"
    run_yaml.font.size = Pt(9)
    run_yaml.font.color.rgb = RGBColor(30, 30, 46)
    # Fondo gris al párrafo
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OXE
    pPr = p_yaml._p.get_or_add_pPr()
    shd = _OXE("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), "E2E8F0")
    pPr.append(shd)

    add_step(4, "Haz commit y push del workflow", [
        "Guarda el archivo sync.yml",
        "En la terminal ejecuta:",
        "  git add .github/workflows/sync.yml",
        "  git commit -m 'Agregar workflow de GitHub Actions'",
        "  git push",
        "GitHub detectara el archivo y programara las ejecuciones automaticamente",
    ])

    add_step(5, "Verifica que funciona", [
        "Ve a tu repositorio en GitHub",
        "Haz clic en la pestana  Actions",
        "Veras el workflow  'Canvas Notion Sync'  listado",
        "Para probarlo manualmente: haz clic en el workflow → 'Run workflow' → 'Run workflow'",
        "Si todo esta bien veras una palomita verde (✓)",
        "Si hay error, haz clic en el job fallido para ver el log completo",
    ])

    doc.add_paragraph()
    add_para(
        "Consejo: ajusta los horarios del cron a tu preferencia. El formato es: "
        "minuto hora dia-mes mes dia-semana (en UTC). Colombia es UTC-5.",
        bold=False, color=(245, 158, 11),
    )

    doc.add_page_break()

    # ── SECCIÓN 6: Resumen ────────────────────────────────────────────────────
    add_heading("6. Resumen de credenciales", level=1)
    doc.add_paragraph("Una vez que tengas todos los datos, el asistente los verificará automáticamente:")

    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    headers = [("Variable", "Descripción")]
    rows_data = [
        ("CANVAS_API_TOKEN", "Token de acceso de Canvas LMS"),
        ("CANVAS_BASE_URL", "URL base de tu Canvas (ej: https://umb.instructure.com/api/v1)"),
        ("NOTION_API_TOKEN", "Token secreto de la integración de Notion"),
        ("NOTION_DATABASE_ID", "ID de la base de datos de Notion"),
        ("TELEGRAM_BOT_TOKEN", "Token del bot de Telegram"),
        ("TELEGRAM_CHAT_ID", "Tu ID de usuario/chat en Telegram"),
    ]
    all_rows = headers + rows_data
    for i, (col1, col2) in enumerate(all_rows):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    add_para(
        "¡Listo! Con el asistente de configuración puedes pegar cada credencial y se "
        "verificará en tiempo real antes de guardar.",
        bold=False,
        color=(34, 197, 94),
    )

    doc.save(str(dest))
    return True, str(dest)


# ══════════════════════════════════════════════════════════════════════════════
# GUI PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

class SetupWizard(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Canvas → Notion  |  Asistente de configuración")
        self.configure(bg=BG)
        self.resizable(False, False)

        # Centrar ventana
        w, h = 780, 620
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        self.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

        self.values = {
            "canvas_url":    tk.StringVar(value="https://umb.instructure.com/api/v1"),
            "canvas_token":  tk.StringVar(),
            "notion_token":  tk.StringVar(),
            "notion_db":     tk.StringVar(),
            "tg_token":      tk.StringVar(),
            "tg_chat":       tk.StringVar(),
        }
        self.valid = {k: False for k in self.values}
        self.current_page = 0

        self._build_ui()
        self.show_page(0)

    # ── UI base ───────────────────────────────────────────────────────────────
    def _build_ui(self):
        # Header
        hdr = tk.Frame(self, bg=ACCENT, height=56)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="⚙  Canvas → Notion Sync  — Asistente de configuración",
                 bg=ACCENT, fg="white", font=("Segoe UI", 12, "bold")).pack(side="left", padx=20)

        # Progress bar strip
        self.progress_frame = tk.Frame(self, bg=BG2, height=6)
        self.progress_frame.pack(fill="x")
        self.progress_bar = tk.Frame(self.progress_frame, bg=ACCENT2, height=6)
        self.progress_bar.place(x=0, y=0, relwidth=0.0, height=6)

        # Content area
        self.content = tk.Frame(self, bg=BG)
        self.content.pack(fill="both", expand=True, padx=30, pady=20)

        # Footer buttons
        footer = tk.Frame(self, bg=BG, height=56)
        footer.pack(fill="x", side="bottom")
        footer.pack_propagate(False)

        self.btn_prev = tk.Button(footer, text="← Anterior", command=self.prev_page,
                                  bg=BG2, fg=FG, font=FONT, relief="flat",
                                  activebackground=ACCENT, activeforeground="white",
                                  padx=16, pady=6)
        self.btn_prev.pack(side="left", padx=20)

        self.btn_next = tk.Button(footer, text="Siguiente →", command=self.next_page,
                                  bg=ACCENT, fg="white", font=FONT_B, relief="flat",
                                  activebackground=ACCENT2, activeforeground="white",
                                  padx=16, pady=6)
        self.btn_next.pack(side="right", padx=20)

        self.page_label = tk.Label(footer, text="", bg=BG, fg=FG2, font=FONT_SM)
        self.page_label.pack(side="right", padx=10)

    # ── Páginas ───────────────────────────────────────────────────────────────
    PAGES = ["Bienvenida", "Canvas LMS", "Notion Token", "Notion DB", "Telegram", "Guardar"]

    def show_page(self, idx: int):
        for w in self.content.winfo_children():
            w.destroy()
        self.current_page = idx
        total = len(self.PAGES)
        self.progress_bar.place(relwidth=(idx + 1) / total)
        self.page_label.config(text=f"Paso {idx+1} de {total}")
        self.btn_prev.config(state="normal" if idx > 0 else "disabled")
        self.btn_next.config(text="Finalizar" if idx == total - 1 else "Siguiente →")

        pages = [
            self._page_welcome,
            self._page_canvas,
            self._page_notion_token,
            self._page_notion_db,
            self._page_telegram,
            self._page_finish,
        ]
        pages[idx]()

    def prev_page(self):
        if self.current_page > 0:
            self.show_page(self.current_page - 1)

    def next_page(self):
        if self.current_page == len(self.PAGES) - 1:
            self._finalize()
        else:
            self.show_page(self.current_page + 1)

    # ── helpers de layout ─────────────────────────────────────────────────────
    def _title(self, text: str):
        tk.Label(self.content, text=text, bg=BG, fg=ACCENT2,
                 font=FONT_H, anchor="w").pack(fill="x", pady=(0, 4))

    def _subtitle(self, text: str):
        tk.Label(self.content, text=text, bg=BG, fg=FG2,
                 font=FONT_SM, anchor="w", wraplength=700,
                 justify="left").pack(fill="x", pady=(0, 12))

    def _field(self, label: str, var: tk.StringVar, show="", placeholder=""):
        frame = tk.Frame(self.content, bg=BG)
        frame.pack(fill="x", pady=4)
        tk.Label(frame, text=label, bg=BG, fg=FG, font=FONT_B, width=22,
                 anchor="w").pack(side="left")
        entry = tk.Entry(frame, textvariable=var, show=show, bg=BG2, fg=FG,
                         font=FONT, relief="flat", insertbackground=FG,
                         width=52)
        entry.pack(side="left", ipady=5, padx=(4, 0))
        if placeholder and not var.get():
            entry.insert(0, placeholder)
            entry.config(fg=FG2)
            def on_focus_in(e, ent=entry, ph=placeholder, v=var):
                if ent.get() == ph:
                    ent.delete(0, "end")
                    ent.config(fg=FG)
            def on_focus_out(e, ent=entry, ph=placeholder, v=var):
                if not ent.get():
                    ent.insert(0, ph)
                    ent.config(fg=FG2)
            entry.bind("<FocusIn>", on_focus_in)
            entry.bind("<FocusOut>", on_focus_out)
        return entry

    def _status_label(self) -> tk.Label:
        lbl = tk.Label(self.content, text="", bg=BG, fg=FG2,
                       font=FONT_SM, anchor="w", wraplength=700)
        lbl.pack(fill="x", pady=(4, 0))
        return lbl

    def _verify_btn(self, text: str, cmd) -> tk.Button:
        btn = tk.Button(self.content, text=text, command=cmd,
                        bg=BG2, fg=ACCENT2, font=FONT, relief="flat",
                        activebackground=ACCENT, activeforeground="white",
                        padx=12, pady=4)
        btn.pack(anchor="w", pady=(6, 0))
        return btn

    def _set_status(self, lbl: tk.Label, ok: bool, msg: str):
        lbl.config(text=("✓  " if ok else "✗  ") + msg,
                   fg=SUCCESS if ok else ERROR)

    # ── Página 0: Bienvenida ──────────────────────────────────────────────────
    def _page_welcome(self):
        self._title("Bienvenido al asistente de configuración")
        self._subtitle(
            "Esta herramienta te guiará para ingresar y verificar todas las credenciales "
            "necesarias para sincronizar Canvas LMS con Notion y recibir notificaciones en Telegram."
        )
        tk.Label(self.content, text="Necesitarás obtener:", bg=BG, fg=FG, font=FONT_B).pack(anchor="w", pady=(10, 4))
        items = [
            "🎓  Token de acceso de Canvas LMS",
            "🗒️  Token secreto de integración de Notion",
            "🗄️  ID de la base de datos de Notion",
            "🤖  Token del bot de Telegram",
            "💬  Tu Chat ID de Telegram",
        ]
        for item in items:
            tk.Label(self.content, text=f"   {item}", bg=BG, fg=FG2, font=FONT,
                     anchor="w").pack(fill="x", pady=2)

        tk.Label(self.content, text="", bg=BG).pack()

        def open_guide():
            dest = Path.home() / "Desktop" / "Guia_Canvas_Notion.docx"
            ok, msg = generate_guide_docx(dest)
            if ok:
                os.startfile(str(dest))
                messagebox.showinfo("Guía generada", f"Guía guardada en:\n{dest}")
            else:
                messagebox.showerror("Error", msg)

        btn_guide = tk.Button(
            self.content,
            text="📄  Descargar guía de instrucciones (Word)",
            command=open_guide,
            bg=BG2, fg=ACCENT2, font=FONT_B, relief="flat",
            activebackground=ACCENT, activeforeground="white",
            padx=16, pady=8,
        )
        btn_guide.pack(anchor="w", pady=(8, 0))
        tk.Label(self.content,
                 text="El archivo Word se guarda en tu Escritorio y se abre automáticamente.",
                 bg=BG, fg=FG2, font=FONT_SM).pack(anchor="w")

    # ── Página 1: Canvas ──────────────────────────────────────────────────────
    def _page_canvas(self):
        self._title("Canvas LMS — Token de acceso")
        self._subtitle(
            "Ve a  Canvas → Configuración → Tokens de acceso aprobados → + Nuevo token.\n"
            "Copia el token generado (solo se muestra una vez)."
        )
        self._field("URL base de Canvas:", self.values["canvas_url"])
        self._field("Token de Canvas:", self.values["canvas_token"], show="•",
                    placeholder="7~AbCd...")
        status = self._status_label()
        if self.valid["canvas_token"]:
            status.config(text="✓  Token verificado", fg=SUCCESS)

        def verify():
            ok, msg = validate_canvas_token(
                self.values["canvas_token"].get().strip(),
                self.values["canvas_url"].get().strip()
            )
            self.valid["canvas_token"] = ok
            self._set_status(status, ok, msg)
        self._verify_btn("Verificar conexión →", verify)

    # ── Página 2: Notion token ────────────────────────────────────────────────
    def _page_notion_token(self):
        self._title("Notion — Token secreto de integración")
        self._subtitle(
            "Ve a  notion.so/my-integrations → Nueva integración → copia el token secreto.\n"
            "El token empieza con  ntn_  o  secret_"
        )
        self._field("Token de Notion:", self.values["notion_token"], show="•",
                    placeholder="ntn_AbCd...")
        status = self._status_label()
        if self.valid["notion_token"]:
            status.config(text="✓  Token verificado", fg=SUCCESS)

        def verify():
            ok, msg = validate_notion_token(self.values["notion_token"].get().strip())
            self.valid["notion_token"] = ok
            self._set_status(status, ok, msg)
        self._verify_btn("Verificar token →", verify)

    # ── Página 3: Notion DB ───────────────────────────────────────────────────
    def _page_notion_db(self):
        self._title("Notion — ID de la base de datos")
        self._subtitle(
            "Abre la base de datos en el navegador. La URL contiene el ID:\n"
            "notion.so/workspace/[ID-DE-32-CHARS]?v=...\n"
            "Recuerda compartir la BD con tu integración (⋯ → Conexiones)."
        )
        self._field("ID de la base de datos:", self.values["notion_db"],
                    placeholder="abcdef12-3456-7890-abcd-ef1234567890")
        status = self._status_label()
        if self.valid["notion_db"]:
            status.config(text="✓  Base de datos verificada", fg=SUCCESS)

        def verify():
            ok, msg = validate_notion_db(
                self.values["notion_token"].get().strip(),
                self.values["notion_db"].get().strip()
            )
            self.valid["notion_db"] = ok
            self._set_status(status, ok, msg)
        self._verify_btn("Verificar base de datos →", verify)

    # ── Página 4: Telegram ────────────────────────────────────────────────────
    def _page_telegram(self):
        self._title("Telegram — Bot y Chat ID")
        self._subtitle(
            "Busca @BotFather en Telegram, escribe /newbot y sigue las instrucciones.\n"
            "Para tu Chat ID busca @userinfobot en Telegram y envíale cualquier mensaje."
        )
        self._field("Token del bot:", self.values["tg_token"], show="•",
                    placeholder="1234567890:AAFb...")
        self._field("Chat ID:", self.values["tg_chat"],
                    placeholder="123456789")
        status = self._status_label()
        if self.valid["tg_token"]:
            status.config(text="✓  Bot verificado", fg=SUCCESS)

        def verify():
            ok, msg = validate_telegram(
                self.values["tg_token"].get().strip(),
                self.values["tg_chat"].get().strip()
            )
            self.valid["tg_token"] = ok
            self._set_status(status, ok, msg)
        self._verify_btn("Enviar mensaje de prueba →", verify)

    # ── Página 5: Guardar ─────────────────────────────────────────────────────
    def _page_finish(self):
        self._title("Resumen y guardar configuración")
        self._subtitle("Revisa que todo esté verificado antes de guardar.")

        checks = [
            ("Canvas URL",    self.values["canvas_url"].get(),   True),
            ("Canvas Token",  "•" * 12,                         self.valid["canvas_token"]),
            ("Notion Token",  "•" * 12,                         self.valid["notion_token"]),
            ("Notion DB ID",  self.values["notion_db"].get(),    self.valid["notion_db"]),
            ("Telegram Bot",  "•" * 12,                         self.valid["tg_token"]),
            ("Telegram Chat", self.values["tg_chat"].get(),      True),
        ]
        for label, val, ok in checks:
            row = tk.Frame(self.content, bg=BG)
            row.pack(fill="x", pady=2)
            color = SUCCESS if ok else WARN
            icon  = "✓" if ok else "⚠"
            tk.Label(row, text=f"{icon}  {label}:", bg=BG, fg=color,
                     font=FONT_B, width=18, anchor="w").pack(side="left")
            tk.Label(row, text=val, bg=BG, fg=FG2, font=FONT).pack(side="left")

        tk.Label(self.content, text="", bg=BG).pack()
        tk.Label(
            self.content,
            text="Al hacer clic en 'Finalizar' se guardará el archivo .env en la carpeta de la aplicación.",
            bg=BG, fg=FG2, font=FONT_SM, wraplength=700, justify="left",
        ).pack(anchor="w")

    def _finalize(self):
        dest = Path(sys.executable).parent / ".env" if hasattr(sys, "_MEIPASS") \
               else Path(__file__).parent.parent / ".env"
        try:
            vals = {k: v.get().strip() for k, v in self.values.items()}
            save_env(dest, vals)
            messagebox.showinfo(
                "¡Configuración guardada!",
                f"Archivo .env guardado en:\n{dest}\n\n"
                "Ahora puedes cerrar este asistente y ejecutar la aplicación principal."
            )
            self.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar:\n{e}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = SetupWizard()
    app.mainloop()
